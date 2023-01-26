import math
import os
import json
import shutil

import matplotlib.pyplot as plt
import pandas as pd

import TBAG
from TBAG import Room, getLongestString
from docx.shared import Pt, RGBColor
from docx import Document
from datetime import datetime

class GlobalRoomReport:
    def __init__(self, room, event, user_count):
        self.room = room
        self.event = event
        self.num_of_success_no_fail = 0
        self.num_of_success_after_fail = 0
        self.num_of_fail_no_success = 0
        self.user_count = user_count

    def increment_data_for_user(self, user_filename):
        data = json.load(open(user_filename))

        if data['tracking'][self.room][self.event]["fail"] == 0 and data['tracking'][self.room][self.event]["success"]:
            self.num_of_success_no_fail += 1
        elif data['tracking'][self.room][self.event]["fail"] > 0 and data['tracking'][self.room][self.event]["success"]:
            self.num_of_success_after_fail += 1
        elif data['tracking'][self.room][self.event]["fail"] > 0 and not data['tracking'][self.room][self.event]["success"]:
            self.num_of_fail_no_success += 1

    def calc_perc_of_fail_no_success(self):
        return round((self.num_of_fail_no_success / self.user_count) * 100, 2)

    def calc_perc_of_success_after_fail(self):
        return round((self.num_of_success_after_fail / self.user_count) * 100, 2)

    def calc_perc_of_success_no_fail(self):
        return round((self.num_of_success_no_fail / self.user_count) * 100, 2)


class UserRoomReport:
    def __init__(self, data_filename):
        self.filename = data_filename

        self.data = json.load(open(self.filename))
        self.map = self.draw_map()
        self.sum_success_no_fail = 0
        self.sum_success_after_fail = 0
        self.sum_fail_no_success = 0

    def get_player_name(self):
        try:
            return self.data['name']
        except KeyError:
            return '<no player name>'

    def get_date_of_access(self):
        try:
            return self.data['date_created']
        except KeyError:
            return 'no date'

    def get_stats_for_event(self, room, event):
        num_of_success_no_fail = 0
        num_of_success_after_fail = 0
        num_of_fail_no_success = 0

        if self.data['tracking'][room][event]["fail"] == 0 and self.data['tracking'][room][event]["success"]:
            num_of_success_no_fail += 1
            self.sum_success_no_fail += 1
        if self.data['tracking'][room][event]["fail"] > 0 and self.data['tracking'][room][event]["success"]:
            num_of_success_after_fail += 1
            self.sum_success_after_fail += 1
        if self.data['tracking'][room][event]["fail"] > 0 and not self.data['tracking'][room][event]["success"]:
            num_of_fail_no_success += 1
            self.sum_fail_no_success += 1

        return num_of_success_no_fail, num_of_success_after_fail, num_of_fail_no_success

    def get_rooms_events(self):
        rooms_events = {}
        for room in self.data['tracking']:
            events = []
            for event in self.data['tracking'][room]:
                event_dict = {}
                event_dict["name"] = self.data['tracking'][room][event]["name"]
                event_dict["id"] = event
                events.append(event_dict)
            rooms_events[room] = events
        return rooms_events

    def room_get_visited(self, roomid=None):
        visited = dict(self.data["visited"])

        # ensure gate is x
        if roomid == "gate":
            return "x"

        if visited[roomid] is True:
            return "x"
        elif visited[roomid] == "current":
            return "x"
        else:
            return ""

    def draw_map(self):
        gate = Room("gate")
        skr = Room("skr")
        adminoffice = Room("adminoffice")
        newsroom = Room("newsroom")
        openoffice = Room("openoffice")
        it = Room("it")
        reception = Room("reception")
        trading = Room("trading")
        youroffice = Room("youroffice")

        trainstation = Room("trainstation")
        home = Room("home")
        map = [
            ['', skr.name, adminoffice.name, newsroom.name, '', trainstation.name, home.name, ''],
            ['', self.room_get_visited(skr.roomid), self.room_get_visited(adminoffice.roomid),
             self.room_get_visited(newsroom.roomid), '---------->', self.room_get_visited(trainstation.roomid),
             self.room_get_visited(home.roomid), ''],
            #
            ['', openoffice.name, youroffice.name, '', '', '', '', ''],
            ['', self.room_get_visited(openoffice.roomid), self.room_get_visited(youroffice.roomid), '', '', '', '',
             ''],
            #
            [gate.name, reception.name, trading.name, '', it.name, '', '', ''],
            [self.room_get_visited(gate.roomid), self.room_get_visited(reception.roomid),
             self.room_get_visited(trading.roomid), '', self.room_get_visited(it.roomid), '', '', '']
        ]

        output = ''
        longeststring = getLongestString(map)
        line = '+---'
        linelength = len(line)
        numofcols = len(map[0])

        FIELDSIZE = len(longeststring)

        rowsize = (FIELDSIZE + 1) * numofcols
        numoflines = math.floor(rowsize / linelength)
        left = rowsize % linelength

        spaces = [None] * numofcols

        for i, a in enumerate(map):
            if i % 2 == 0:
                output += '\n' + line * numoflines + '-' * left + '+'
            for j, b in enumerate(a):
                if i % 2 == 0:
                    # getting + spacing - get length of roomname
                    #spaces[j] = len(b)

                    # printing
                    if j == 0:
                        output += '\n|{:^{x}}'.format(b, x=FIELDSIZE)
                    else:
                        output += '|{:^{x}}'.format(b, x=FIELDSIZE)
                else:
                    if j == 0:
                        output += '\n|{:^{x}}'.format(b, x=FIELDSIZE)
                    else:
                        output += '|{:^{x}}'.format(b, x=FIELDSIZE)
            output += '|'
        output += '\n' + line * numoflines + '-' * left + '+\n\n'
        return output


class ReportGenerator:
    def __init__(self, report_dir_path, userdata_dir_path, rooms_dir_path):
        self.__date = datetime.today().strftime('%Y-%m-%d')

        self.__report_dir_path = report_dir_path
        self.__userdata_dir_path = userdata_dir_path
        self.__rooms_dir_path = rooms_dir_path
        self.__archive_dir_path = os.path.join("reports", self.__date, "archive")
        self.__number_of_rooms = self.get_number_of_rooms()
        self.__number_of_participants = self.get_number_of_participants()
    def generate_reports_individual(self):
        for f in os.listdir(self.__userdata_dir_path):
            if f.endswith(".json"):
                datafile_path = os.path.join(self.__userdata_dir_path, f)
                self.generate_single_report(datafile_path=datafile_path, datafile_name=f)
        print("Reports successfully created!")

    def generate_single_report(self, datafile_path, datafile_name):
        user_report = UserRoomReport(data_filename=datafile_path)
        document = Document()
        player_name = user_report.get_player_name()
        date_of_access = user_report.get_date_of_access()

        document.add_heading('Performance report', 0)
        p = document.add_paragraph('Performance report for player: ')
        p.add_run(player_name+"\n").bold = True
        p.add_run(f"The {player_name} first accessed the game at: {date_of_access}")

        # insert map
        map_p = document.add_paragraph()
        map_p.alignment = 1
        map_p_run = map_p.add_run(user_report.map)
        map_p_run.font.name = 'Courier New'
        map_p_run.font.size = Pt(5)

        document.add_heading('Introduction', 1)
        document.add_paragraph('The game included various security requirements testing in different rooms.')

        rooms_events = user_report.get_rooms_events()
        for room in rooms_events:
            roomname = TBAG.Room(room).name
            document.add_heading('Section ' + roomname, 2)


            for event in rooms_events[room]:
                eventid = event["id"]
                eventname = event["name"]

                #document.add_heading(eventname, 4)

                event_p = document.add_paragraph()
                event_p = event_p.add_run("- "+eventname)
                event_p.bold = True
                event_p.font.name = 'Cambria'
                event_p.font.size = Pt(11)
                event_p.font.color.rgb = RGBColor(0x00, 0x80, 0x80)

                num_of_success_no_fail, num_of_success_after_fail, num_of_fail_no_success \
                    = user_report.get_stats_for_event(room, eventid)

                document.add_paragraph('Number of success with no failure: ' +
                                       str(num_of_success_no_fail))
                document.add_paragraph('Number of success after failure: ' +
                                       str(num_of_success_after_fail))
                document.add_paragraph('Number of failure with no success: ' +
                                       str(num_of_fail_no_success))

        document.add_heading('Summing up', 1)
        document.add_paragraph('Number of success with no failure in all rooms: ' +
                               str(user_report.sum_success_no_fail))
        document.add_paragraph('Number of success after failure in all rooms: ' +
                               str(user_report.sum_success_after_fail))
        document.add_paragraph('Number of failure with no success in all rooms: ' +
                               str(user_report.sum_fail_no_success))

        report_path_indiv = os.path.join(self.__report_dir_path, self.__date)
        if not os.path.exists(report_path_indiv):
            os.makedirs(report_path_indiv)

        report_path = os.path.join(report_path_indiv, "report_" + player_name + "_" + datafile_name[0:5] + '.docx')
        document.save(report_path)
        print(f"Report for user {player_name} successfully created! You can find it at: {report_path}")

    def generate_report_all(self):
        userfile = os.path.join(self.__userdata_dir_path,os.listdir(self.__userdata_dir_path)[0])
        user_report = UserRoomReport(data_filename=userfile)
        rooms_events = user_report.get_rooms_events()

        document = Document()
        print("Creating global report...")
        document.add_heading('Global performance report', 0)

        document.add_heading('Introduction', 1)
        document.add_paragraph('The game included various security requirements testing in different rooms.')
        document.add_paragraph('Number of rooms: ' +
                               str(self.__number_of_rooms))
        document.add_paragraph('Number of participants: ' +
                               str(self.__number_of_participants))

        for room in rooms_events:
            roomname = TBAG.Room(room).name
            document.add_heading('Section ' + roomname, 2)

            for event in rooms_events[room]:
                eventid = event["id"]
                eventname = event["name"]

                event_p = document.add_paragraph()
                event_p = event_p.add_run("- " + eventname)
                event_p.bold = True
                event_p.font.name = 'Cambria'
                event_p.font.size = Pt(11)
                event_p.font.color.rgb = RGBColor(0x00, 0x80, 0x80)

                global_report = GlobalRoomReport(room, eventid, user_count=self.__number_of_participants)
                for f in os.listdir(self.__userdata_dir_path):
                    if f.endswith(".json"):
                        datafile_path = os.path.join(self.__userdata_dir_path, f)
                        global_report.increment_data_for_user(datafile_path)

                perc_of_success_no_fail = global_report.calc_perc_of_success_no_fail()
                perc_of_success_after_fail = global_report.calc_perc_of_success_after_fail()
                perc_of_fail_no_success = global_report.calc_perc_of_fail_no_success()


                document.add_paragraph('Number and percentage of users who succeeded with no failure: ' +
                                       str(global_report.num_of_success_no_fail)+" | "+str(perc_of_success_no_fail)+"%")

                document.add_paragraph('Number and percentage of users who succeeded after failure: ' +
                                       str(global_report.num_of_success_after_fail)+" | "+str(perc_of_success_after_fail)+"%")

                document.add_paragraph('Number and percentage of users who failed with no success: ' +
                                       str(global_report.num_of_fail_no_success)+" | "+str(perc_of_fail_no_success)+"%")

                # Add piechart
                plt.rcParams["figure.figsize"] = [5.00, 3.50]
                plt.rcParams["figure.autolayout"] = True
                left = global_report.user_count - global_report.num_of_success_no_fail - global_report.num_of_success_after_fail - global_report.num_of_fail_no_success

                df = pd.DataFrame(
                    {
                        'labels': ['Succeeded', 'Succeeded after failure', 'Failed', 'Did not engage'],
                        '': [global_report.num_of_success_no_fail, global_report.num_of_success_after_fail,
                             global_report.num_of_fail_no_success, left]
                    }
                )
                df.set_index('labels').plot.pie(y='', legend=False,
                                                autopct=lambda p: format(p, '.2f') + "%" if p > 0 else None)
                plt.savefig('piechart.png')
                document.add_picture('piechart.png')
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = 0
                os.remove('piechart.png')

        if not os.path.exists(self.__report_dir_path,):
            os.makedirs(self.__report_dir_path)

        report_path = os.path.join(self.__report_dir_path, "globalreport_" + self.__date + '.docx')
        document.save(report_path)
        print(f"Global report successfully created! You can find it at: {report_path}")

    def move_reports(self):

        if not os.path.exists(self.__archive_dir_path ):
            os.makedirs(self.__archive_dir_path )

        try:
            for f in os.listdir(self.__userdata_dir_path):
                if f.endswith(".json"):
                    datafile_path = os.path.join(self.__userdata_dir_path, f)
                    shutil.move(datafile_path, self.__archive_dir_path )
        except FileNotFoundError:
            print(f"The {self.__userdata_dir_path} folder doesn't exist. Please make the script is in "
                  f"the game's backend directory.")
            exit()

        print("User data files successfully moved to "+self.__archive_dir_path)

    def get_number_of_rooms(self):
        try:
            num_of_rooms = 0
            for r_file in os.listdir(self.__rooms_dir_path):
                if r_file.endswith('.yaml'):
                    num_of_rooms += 1
            return num_of_rooms
        except FileNotFoundError:
            print(f"The {self.__rooms_dir_path} folder doesn't exist. Please make the script is in "
                  f"the game's backend directory.")
            exit()
    def get_number_of_participants(self):
        try:
            num_of_participants = 0
            for u_file in os.listdir(self.__userdata_dir_path):
                if u_file.endswith('.json'):
                    num_of_participants += 1
            return num_of_participants
        except FileNotFoundError:
            print(f"The {self.__userdata_dir_path} folder doesn't exist. Please make the script is in "
                  f"the game's backend directory.")
            exit()

if __name__ == '__main__':
    report_directory_path = "reports"
    userdata_directory = "persistent"
    rooms_directory = "rooms"

    report_generator = ReportGenerator(report_dir_path=report_directory_path,
                                       userdata_dir_path=userdata_directory,
                                       rooms_dir_path=rooms_directory)

    while True:
        print("\n1. Generate a report for every player")
        print("2. Generate single report for all players")
        print("3. Move userdata files to archive")
        print("Enter 'q' to quit\n")

        choice = input("Enter your choice: ")

        if choice == '1':
            report_generator.generate_reports_individual()
            input("Press enter to continue..")
        elif choice == '2':
            report_generator.generate_report_all()
            input("Press enter to continue..")
        elif choice == '3':
            report_generator.move_reports()
            input("Press enter to continue..")
        elif choice == 'q':
            break
        else:
            print("Invalid input, please try again")